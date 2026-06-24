# -*- coding: utf-8 -*-
"""20_table1_data_records.py — Table 1 数据记录三线表
真实样本：6 样本(MOCK_1-3/PR8_1-3)；文件类型来自 anchored_facts。GEO 占位。
三线表：顶粗线 + 表头下细线 + 底粗线，无竖线无内部横线。
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\IAV_m6A\SciData_descriptor\submission\06_tables\Table1_data_records.docx"
os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()
# title
p = doc.add_paragraph()
r = p.add_run("Table 1. Data records: samples, repository, accession, and file types.")
r.bold = True; r.font.size = Pt(10); r.font.name = "Times New Roman"

rows_data = [
    ["Sample", "Condition", "Repository", "Accession", "Files"],
    ["MOCK_1", "Mock-infected, rep 1", "GEO", "GSM[TBD]", "MOCK_1.txt"],
    ["MOCK_2", "Mock-infected, rep 2", "GEO", "GSM[TBD]", "MOCK_2.txt"],
    ["MOCK_3", "Mock-infected, rep 3", "GEO", "GSM[TBD]", "MOCK_3.txt"],
    ["PR8_1", "IAV PR8, MOI 0.5, 24 h, rep 1", "GEO", "GSM[TBD]", "PR8_1.txt"],
    ["PR8_2", "IAV PR8, MOI 0.5, 24 h, rep 2", "GEO", "GSM[TBD]", "PR8_2.txt"],
    ["PR8_3", "IAV PR8, MOI 0.5, 24 h, rep 3", "GEO", "GSM[TBD]", "PR8_3.txt"],
    ["All 6", "Series (platform GPL25759)", "GEO", "GSE[TBD]", "Matrix.xlsx (40,956 probes); processed profile tables"],
]

table = doc.add_table(rows=len(rows_data), cols=5)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, row in enumerate(rows_data):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = ""
        cp = cell.paragraphs[0]
        cr = cp.add_run(val)
        cr.font.size = Pt(9); cr.font.name = "Times New Roman"
        if i == 0: cr.bold = True

def set_border(cell, edge, sz, val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = tcPr.find(qn('w:tcBorders'))
    if tcB is None:
        tcB = OxmlElement('w:tcBorders'); tcPr.append(tcB)
    el = tcB.find(qn('w:'+edge))
    if el is None:
        el = OxmlElement('w:'+edge); tcB.append(el)
    el.set(qn('w:val'), val); el.set(qn('w:sz'), str(sz))
    el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')

ncols = 5; nrows = len(rows_data)
for j in range(ncols):
    set_border(table.cell(0, j), 'top', 18)        # top thick
    set_border(table.cell(0, j), 'bottom', 8)      # header thin
    set_border(table.cell(nrows-1, j), 'bottom', 18)  # bottom thick

# footnote
fn = doc.add_paragraph()
fr = fn.add_run("Each sample is a single Arraystar Human mRNA&lncRNA Epitranscriptomic "
                "Microarray (8×60K) hybridisation. Raw Agilent Feature Extraction text "
                "files (.txt), the normalised probe intensity matrix, and processed m6A "
                "quantity / expression profile tables are deposited under the GEO Series. "
                "Accession numbers will be provided on data release.")
fr.font.size = Pt(8); fr.font.name = "Times New Roman"; fr.italic = True

doc.save(OUT)
print("SAVED:", OUT)
