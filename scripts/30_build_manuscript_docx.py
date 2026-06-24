# -*- coding: utf-8 -*-
"""30_build_manuscript_docx.py — 构建投稿用 manuscript.docx
1) 读 md，剥离 HTML 工作备注注释（docx 不带工作备注）
2) pypandoc 转 docx
3) 后处理：[INSERT ... ABOUT HERE] 占位 → 14pt 粗体居中
复用 TBI SciData descriptor 的构建方法(23_finalize_manuscript.py)，占位格式适配本项目。
"""
import re, os
from pathlib import Path
import pypandoc
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

MS = Path(r"D:\IAV_m6A\SciData_descriptor\manuscript")
SRC = MS / "data_descriptor_submission.md"
CLEAN = MS / "_build_clean.md"
OUT = MS / "data_descriptor_submission.docx"
SUB = Path(r"D:\IAV_m6A\SciData_descriptor\submission\01_manuscript\manuscript.docx")

# 1) strip HTML comments (工作备注)
text = SRC.read_text(encoding="utf-8")
n_comments = text.count("<!--")
text_clean = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
# remove standalone '---' horizontal-rule lines (avoid pandoc YAML front-matter misparse;
# SciData docx 不需要水平分隔线)
text_clean = re.sub(r"(?m)^\s*---\s*$", "", text_clean)
# collapse >2 blank lines
text_clean = re.sub(r"\n{3,}", "\n\n", text_clean)
CLEAN.write_text(text_clean, encoding="utf-8")
print(f"stripped comments: {n_comments} -> {text_clean.count('<!--')}; removed --- rules")

# 2) convert
pypandoc.convert_file(str(CLEAN), "docx", outputfile=str(OUT),
    extra_args=["--from=markdown+pipe_tables+fenced_code_blocks"])
print("converted:", OUT.name)

# 3) post-process placeholders
PH = re.compile(r"^\[INSERT (FIGURE \d|TABLE \d|SUPPLEMENTARY FIGURE \d) ABOUT HERE\]$")
doc = Document(str(OUT))
n = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    m = PH.match(txt)
    if m:
        label = m.group(1).title().replace("Supplementary Figure","Supplementary Figure")
        for run in p.runs: run.text = ""
        r = p.add_run(label)
        r.bold = True; r.font.size = Pt(14); r.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = doc.styles["Normal"]
        n += 1
doc.save(str(OUT))
print(f"placeholders styled: {n}")

# 4) copy to submission
import shutil
SUB.parent.mkdir(parents=True, exist_ok=True)
shutil.copy(str(OUT), str(SUB))
os.remove(str(CLEAN))
print(f"saved: {OUT}  ({OUT.stat().st_size/1024:.1f} KB)")
print(f"copied to: {SUB}")
