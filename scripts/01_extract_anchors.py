# -*- coding: utf-8 -*-
"""01_extract_anchors.py — 从 IAV_m6A 商业报告抽取真实数据锚点
数据真实性铁律：所有数字必须来自实际文件读取，禁止凭记忆补全。
"""
import os, sys
from docx import Document
import openpyxl

ROOT = r"D:\IAV_m6A\Report\Report"
OUT = r"D:\IAV_m6A\SciData_descriptor\results\anchor_extraction_raw.txt"
os.makedirs(os.path.dirname(OUT), exist_ok=True)

lines = []
def log(s=""):
    lines.append(str(s))

# ===== 1) 主报告 docx 全文段落 =====
docx_path = os.path.join(ROOT, "Report", "Human m6A-mRNA&lncRNA Epitranscriptomic Microarray Report.docx")
log("="*70)
log("[1] MAIN REPORT DOCX: " + docx_path)
log("exists=" + str(os.path.exists(docx_path)))
if os.path.exists(docx_path):
    doc = Document(docx_path)
    log(f"  paragraphs={len(doc.paragraphs)}  tables={len(doc.tables)}")
    log("--- non-empty paragraphs ---")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            log(f"P{i}: {t}")
    log("--- tables ---")
    for ti, tbl in enumerate(doc.tables):
        log(f"=== TABLE {ti} ({len(tbl.rows)}x{len(tbl.columns)}) ===")
        for r in tbl.rows:
            cells = [c.text.strip() for c in r.cells]
            log(" | ".join(cells))

with open(OUT, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

# ===== 2) Key xlsx files: dims + header + sample columns =====
RP = os.path.join(ROOT, "Report")
xlsx_targets = {
    "DEG_diffexpressed": os.path.join(RP, "PaperData", "Differentially Expressed Genes", "mRNA_ExpressionLevel(diffexpressed).xlsx"),
    "DEG_all": os.path.join(RP, "PaperData", "Differentially Expressed Genes", "mRNA_ExpressionLevel(allcomparison).xlsx"),
    "DM_mRNA_diff": os.path.join(RP, "PaperData", "Differentially Methylated Genes", "mRNA_Quantity(diffmethylated).xlsx"),
    "DM_mRNA_all": os.path.join(RP, "PaperData", "Differentially Methylated Genes", "mRNA_Quantity(allcomparison).xlsx"),
    "DM_lncRNA_diff": os.path.join(RP, "PaperData", "Differentially Methylated Genes", "lncRNA_Quantity(diffmethylated).xlsx"),
    "DM_lncRNA_all": os.path.join(RP, "PaperData", "Differentially Methylated Genes", "lncRNA_Quantity(allcomparison).xlsx"),
    "DM_others_diff": os.path.join(RP, "PaperData", "Differentially Methylated Genes", "others_Quantity(diffmethylated).xlsx"),
    "DM_others_all": os.path.join(RP, "PaperData", "Differentially Methylated Genes", "others_Quantity(allcomparison).xlsx"),
    "RawMatrix": os.path.join(RP, "PaperData", "Raw Data Files", "Matrix.xlsx"),
    "mRNA_expr_flagpass": os.path.join(RP, "GeneralData", "Expression Profiling", "mRNA_ExpressionLevel(flag_pass).xlsx"),
    "mRNA_quant_flagpass": os.path.join(RP, "GeneralData", "Methylation Profiling", "mRNA_Quantity(flag_pass).xlsx"),
    "lncRNA_quant_flagpass": os.path.join(RP, "GeneralData", "Methylation Profiling", "lncRNA_Quantity(flag_pass).xlsx"),
    "others_quant_flagpass": os.path.join(RP, "GeneralData", "Methylation Profiling", "others_Quantity(flag_pass).xlsx"),
}
xlsx_targets["m6A_enzymes"] = os.path.join(ROOT, "A549-PR8的m6A数据的再分析 - 2022-4-13", "m6A相关的酶", "m6A酶表达情况xlsx.xlsx")

xlines = []
def xl(s=""):
    xlines.append(str(s))

for name, path in xlsx_targets.items():
    xl("="*70)
    xl(f"[XLSX] {name}")
    xl("path=" + path)
    xl("exists=" + str(os.path.exists(path)))
    if not os.path.exists(path):
        continue
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        for ws in wb.worksheets:
            xl(f"  sheet='{ws.title}'  dims={ws.max_row}x{ws.max_column}")
            # header + first 3 data rows
            for ri, row in enumerate(ws.iter_rows(min_row=1, max_row=4, values_only=True)):
                xl(f"    R{ri}: " + " | ".join("" if v is None else str(v) for v in row))
        wb.close()
    except Exception as e:
        xl("  ERROR reading: " + repr(e))

OUT2 = r"D:\IAV_m6A\SciData_descriptor\results\anchor_xlsx_dims.txt"
with open(OUT2, "w", encoding="utf-8") as f:
    f.write("\n".join(xlines))

